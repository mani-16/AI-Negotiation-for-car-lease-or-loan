"""
Universal document text extraction service.

Supports: PDF, DOCX, XLSX, CSV, TXT, RTF, HTML, EML, MSG,
          JPG, PNG, TIFF, HEIC, BMP, WEBP
Max file size: 50MB
Blocked: .exe .sh .bat .js .py .php

STRATEGY:
  PDF  -> pdfplumber first (native text)
          If text < 100 chars -> fallback to Google Vision OCR
          If Vision fails    -> fallback to pytesseract OCR
          If all fail        -> ask user to re-upload
  DOCX -> python-docx
  XLSX -> openpyxl (all sheets)
  Images (JPG, PNG, TIFF, HEIC, BMP, WEBP) -> Google Vision OCR
          If Vision fails -> fallback to pytesseract OCR
          If both fail    -> ask user to re-upload
  TXT/CSV/HTML/RTF/EML -> direct decode
  MSG  -> extract_msg library
  Unsupported -> raise ValueError

Google Vision OCR:
  - Only called when needed (PDF fallback + images)
  - Credentials from GOOGLE_CLOUD_CREDENTIALS_JSON env var
  - If credentials are empty {} -> skip OCR, return extracted text only
  - Never raise on Vision failure -> log warning, return what we have
"""

import io
import json
import base64
import logging
from pathlib import Path
import pdfplumber
import docx
import openpyxl
from app.core.config import settings
from app.services.ocr_service import ocr_image_bytes, ocr_pdf_bytes

logger = logging.getLogger(__name__)

# Register HEIC support so Pillow can open .heic images
try:
    from pillow_heif import register_heif_opener
    register_heif_opener()
except ImportError:
    logger.info("pillow-heif not installed, HEIC support unavailable")

BLOCKED_EXTENSIONS = {
    ".exe", ".sh", ".bat", ".js", ".py", ".php",
    ".rb", ".pl", ".cmd", ".ps1", ".vbs"
}
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB


class ExtractionService:

    # ─── ENTRY POINT ──────────────────────────────────────────────────
    async def extract_text(
        self,
        file_bytes: bytes,
        filename: str,
        mime_type: str
    ) -> str:
        """
        Main entry point.
        Returns extracted text string.
        Raises ValueError for blocked/unsupported files.
        Raises RuntimeError for extraction failures.
        """
        # Size check
        if len(file_bytes) > MAX_FILE_SIZE:
            raise ValueError(
                f"File too large: {len(file_bytes) // (1024*1024)}MB. "
                f"Maximum allowed: 50MB."
            )

        # Extension check
        ext = Path(filename).suffix.lower()
        if ext in BLOCKED_EXTENSIONS:
            raise ValueError(
                f"File type '{ext}' is not allowed for security reasons."
            )

        # Route by mime type or extension
        if mime_type == "application/pdf" or ext == ".pdf":
            return await self._extract_pdf(file_bytes)

        elif mime_type in (
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document",
            "application/msword"
        ) or ext in (".docx", ".doc"):
            if ext == ".doc" and mime_type == "application/msword":
                # Old .doc format — try text decode as fallback
                try:
                    return self._extract_text_file(file_bytes, ext=ext)
                except Exception:
                    pass
            return self._extract_docx(file_bytes)

        elif mime_type in (
            "application/vnd.openxmlformats-officedocument"
            ".spreadsheetml.sheet",
            "application/vnd.ms-excel"
        ) or ext in (".xlsx", ".xls"):
            return self._extract_xlsx(file_bytes)

        elif mime_type in (
            "image/jpeg", "image/png", "image/tiff",
            "image/heic", "image/bmp", "image/webp"
        ) or ext in (
            ".jpg", ".jpeg", ".png", ".tiff", ".tif",
            ".heic", ".bmp", ".webp"
        ):
            return await self._extract_image_ocr(file_bytes)

        elif mime_type in (
            "text/plain", "text/csv", "text/html",
            "text/rtf", "message/rfc822"
        ) or ext in (
            ".txt", ".csv", ".html", ".htm", ".rtf", ".eml"
        ):
            return self._extract_text_file(file_bytes, ext=ext)

        elif ext == ".msg":
            return self._extract_msg(file_bytes)

        else:
            # Try plain text as last resort
            try:
                return file_bytes.decode("utf-8", errors="replace")
            except Exception:
                raise ValueError(
                    f"Unsupported file type: {mime_type} ({ext})"
                )

    # ─── PDF ──────────────────────────────────────────────────────────
    async def _extract_pdf(self, file_bytes: bytes) -> str:
        """
        Try pdfplumber first for native PDF text.
        If text < 100 chars (scanned PDF), fall back to Google Vision OCR.
        """
        text = ""
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        parts.append(page_text.strip())
                text = "\n\n".join(parts)
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}")
            text = ""

        if len(text.strip()) >= 50:
            return text

        ocr_text = ""
        tess_text = ""

        # Fallback 1: Google Vision OCR for scanned PDFs
        logger.info(
            "PDF text too short, attempting Google Vision OCR fallback"
        )
        try:
            ocr_text = await self._vision_ocr(file_bytes, "application/pdf")
            if ocr_text and len(ocr_text.strip()) >= 50:
                return ocr_text
        except Exception as e:
            logger.warning(f"Google Vision OCR fallback failed: {e}")

        # Fallback 2: pytesseract OCR for scanned PDFs
        logger.info(
            "Google Vision failed or insufficient, "
            "attempting pytesseract OCR fallback"
        )
        try:
            tess_text = ocr_pdf_bytes(file_bytes)
            if tess_text and len(tess_text.strip()) >= 50:
                return tess_text
        except Exception as e:
            logger.warning(f"pytesseract PDF OCR fallback failed: {e}")

        # Return whatever partial text we got from any method
        best = max(
            [text, ocr_text, tess_text],
            key=lambda t: len(t.strip()) if t else 0,
        )
        if best and best.strip():
            return best

        raise RuntimeError(
            "Could not extract text from this PDF. "
            "Both Google Vision and pytesseract OCR failed. "
            "Please re-upload the document in a different format "
            "(e.g. a clearer scan or a native PDF)."
        )

    # ─── DOCX ─────────────────────────────────────────────────────────
    def _extract_docx(self, file_bytes: bytes) -> str:
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        paragraphs.append(" | ".join(cells))
            text = "\n".join(paragraphs)
            if not text.strip():
                raise RuntimeError("No text found in Word document.")
            return text
        except Exception as e:
            raise RuntimeError(f"Word document extraction failed: {e}")

    # ─── XLSX ─────────────────────────────────────────────────────────
    def _extract_xlsx(self, file_bytes: bytes) -> str:
        try:
            wb = openpyxl.load_workbook(
                io.BytesIO(file_bytes), read_only=True, data_only=True
            )
            parts = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                sheet_parts = [f"[Sheet: {sheet_name}]"]
                for row in ws.iter_rows(values_only=True):
                    cells = [
                        str(c) for c in row
                        if c is not None and str(c).strip()
                    ]
                    if cells:
                        sheet_parts.append(" | ".join(cells))
                if len(sheet_parts) > 1:
                    parts.append("\n".join(sheet_parts))
            text = "\n\n".join(parts)
            if not text.strip():
                raise RuntimeError("No data found in spreadsheet.")
            return text
        except Exception as e:
            raise RuntimeError(f"Spreadsheet extraction failed: {e}")

    # ─── IMAGE OCR ────────────────────────────────────────────────────
    async def _extract_image_ocr(self, file_bytes: bytes) -> str:
        """
        Extract text from image.
        Strategy: Google Vision OCR -> pytesseract fallback.
        If both fail, ask user to re-upload.
        """
        # Attempt 1: Google Vision OCR
        vision_text = ""
        try:
            vision_text = await self._vision_ocr(file_bytes, None)
            if vision_text and vision_text.strip():
                return vision_text
        except Exception as e:
            logger.warning(f"Google Vision image OCR failed: {e}")

        # Attempt 2: pytesseract fallback
        logger.info(
            "Google Vision failed or empty, "
            "attempting pytesseract OCR fallback for image"
        )
        try:
            tess_text = ocr_image_bytes(file_bytes)
            if tess_text and tess_text.strip():
                return tess_text
        except Exception as e:
            logger.warning(f"pytesseract image OCR fallback failed: {e}")

        raise RuntimeError(
            "No text could be extracted from this image. "
            "Both Google Vision and pytesseract OCR failed. "
            "Please re-upload a clearer image or convert the "
            "document to a different format (e.g. PDF or DOCX)."
        )

    # ─── TEXT / CSV / HTML / RTF / EML FILES ──────────────────────────
    def _extract_text_file(self, file_bytes: bytes, ext: str = "") -> str:
        try:
            text = file_bytes.decode("utf-8", errors="replace")

            # RTF: strip formatting with striprtf
            if ext == ".rtf" or text.strip().startswith("{\\rtf"):
                try:
                    from striprtf.striprtf import rtf_to_text
                    text = rtf_to_text(text)
                except Exception as e:
                    logger.warning(f"RTF stripping failed, using raw: {e}")

            # HTML: extract visible text with BeautifulSoup
            if ext in (".html", ".htm") or (
                text.strip().startswith("<") and "<html" in text[:500].lower()
            ):
                try:
                    from bs4 import BeautifulSoup
                    soup = BeautifulSoup(text, "html.parser")
                    # Remove script and style elements
                    for tag in soup(["script", "style"]):
                        tag.decompose()
                    text = soup.get_text(separator="\n", strip=True)
                except Exception as e:
                    logger.warning(f"HTML parsing failed, using raw: {e}")

            if not text.strip():
                raise RuntimeError("File appears to be empty.")
            return text
        except RuntimeError:
            raise
        except Exception as e:
            raise RuntimeError(f"Text file extraction failed: {e}")

    # ─── MSG ──────────────────────────────────────────────────────────
    def _extract_msg(self, file_bytes: bytes) -> str:
        try:
            import extract_msg
            import tempfile, os
            with tempfile.NamedTemporaryFile(
                suffix=".msg", delete=False
            ) as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name
            try:
                msg = extract_msg.Message(tmp_path)
                parts = []
                if msg.subject:
                    parts.append(f"Subject: {msg.subject}")
                if msg.sender:
                    parts.append(f"From: {msg.sender}")
                if msg.body:
                    parts.append(msg.body)
                return "\n".join(parts)
            finally:
                os.unlink(tmp_path)
        except ImportError:
            raise RuntimeError(
                "MSG file support not installed. "
                "Please convert to PDF or TXT."
            )
        except Exception as e:
            raise RuntimeError(f"MSG extraction failed: {e}")

    # ─── GOOGLE VISION OCR ────────────────────────────────────────────
    async def _vision_ocr(
        self, file_bytes: bytes, mime_type: str | None
    ) -> str:
        """
        Call Google Cloud Vision OCR.
        Returns empty string if credentials not configured.
        """
        credentials_json = settings.GOOGLE_CLOUD_CREDENTIALS_JSON
        if not credentials_json or credentials_json.strip() in ("{}", ""):
            logger.info(
                "Google Cloud credentials not configured, skipping OCR"
            )
            return ""

        try:
            from google.cloud import vision
            from google.oauth2 import service_account

            creds_dict = json.loads(credentials_json)
            credentials = (
                service_account.Credentials.from_service_account_info(
                    creds_dict,
                    scopes=[
                        "https://www.googleapis.com/auth/cloud-platform"
                    ],
                )
            )
            client = vision.ImageAnnotatorClient(
                credentials=credentials
            )

            if mime_type == "application/pdf":
                image = vision.Image(content=file_bytes)
                response = client.document_text_detection(image=image)
            else:
                image = vision.Image(content=file_bytes)
                response = client.document_text_detection(image=image)

            if response.error.message:
                raise RuntimeError(
                    f"Vision API error: {response.error.message}"
                )

            return response.full_text_annotation.text or ""

        except ImportError:
            logger.warning(
                "google-cloud-vision not installed, OCR unavailable"
            )
            return ""
        except Exception as e:
            logger.warning(f"Google Vision OCR failed: {e}")
            return ""


# Singleton
extraction_service = ExtractionService()
